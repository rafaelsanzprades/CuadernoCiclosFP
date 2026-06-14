"""
generador_pd.py
Genera la Programación Didáctica en .docx y .pdf
Formato: Arial 11, espaciado 1.15, encabezado/pie estilo GM-1-0237
Índice automático referenciado con estilos Word nativos (Heading 1/2 → TOC)

Uso:  python generador_pd.py [module_id] [output_dir]
      python generador_pd.py  (usa datos demo)
"""
import sys, json, os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── helpers ────────────────────────────────────────────────────────────────

def set_font(run, bold=False, size=11, color=None):
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    # Forzar Arial para idiomas latinos
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    rPr.insert(0, rFonts)

def set_para_spacing(para, before=0, after=6, line=1.15):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(11 * line)

def add_bottom_border(para):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)

def add_heading(doc, text, level=1):
    """Añade encabezado con estilo Word Heading 1/2/3 (para TOC automático)"""
    style_name = f"Heading {level}"
    para = doc.add_paragraph(style=style_name)
    run = para.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    if level == 1:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.underline = True
    elif level == 2:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif level == 3:
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run.italic = True

    set_para_spacing(para, before=12, after=4)
    add_bottom_border(para)
    return para

def add_body(doc, text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    set_font(run)
    set_para_spacing(para)
    return para

def add_bullet(doc, text, level=0):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1 + 0.5 * level)
    run = para.add_run(text)
    set_font(run)
    set_para_spacing(para, before=3, after=0)
    return para

def add_toc(doc):
    """Inserta campo TOC de Word (se actualiza al abrir en Word con F9)"""
    para = doc.add_paragraph()
    run = para.add_run()
    fldChar = OxmlElement("w:fldChar")
    fldChar.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run2 = para.add_run()
    run2._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    run3 = para.add_run()
    run3._r.append(fldChar2)

    para2 = doc.add_paragraph()
    run4 = para2.add_run()
    set_font(run4)
    run4._r.append(OxmlElement("w:rPr"))

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")
    run5 = para2.add_run()
    run5._r.append(fldChar3)

    note = doc.add_paragraph()
    rn = note.add_run("⚠ ÍNDICE: Para generar o actualizar el índice automáticamente si añades o modificas títulos, haz clic derecho sobre este bloque y selecciona 'Actualizar campos'.")
    set_font(rn, size=10, bold=True)
    rn.font.color.rgb = RGBColor(0, 0, 0)
    set_para_spacing(note, before=2, after=8)


def setup_header_footer(doc, modulo_nombre, ciclo_nombre, curso_academico):
    """Encabezado y pie idénticos al PDF de referencia"""
    for section in doc.sections:
        section.top_margin    = Cm(1)
        section.bottom_margin = Cm(1)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(1)
        section.different_first_page_header_footer = True

        # ── ENCABEZADO ─────────────────────────────────────────
        header = section.header
        header.is_linked_to_previous = False
        # Tabla 3 columnas: [Ciclo] | [Módulo] | [Curso / Página]
        tbl = header.add_table(2, 3, Cm(18))
        tbl.autofit = False
        tbl.allow_autofit = False
        tbl.style = "Table Grid"
        
        header_widths = [Cm(6.5), Cm(8.5), Cm(3)]
        for row in tbl.rows:
            for idx, width in enumerate(header_widths):
                row.cells[idx].width = width

        labels = ["CICLO FORMATIVO", "MÓDULO PROFESIONAL", "CURSO ACADÉMICO"]
        values = [ciclo_nombre, modulo_nombre, curso_academico]
        for col_idx, (lbl, val) in enumerate(zip(labels, values)):
            # Fila 0: etiqueta
            cell0 = tbl.rows[0].cells[col_idx]
            p0 = cell0.paragraphs[0]
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r0 = p0.add_run(lbl)
            set_font(r0, bold=True, size=8)
            # Fila 1: valor
            cell1 = tbl.rows[1].cells[col_idx]
            p1 = cell1.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = p1.add_run(val)
            set_font(r1, size=8)

        # ── PIE DE PÁGINA ───────────────────────────────────────
        footer = section.footer
        footer.is_linked_to_previous = False
        tbl_f = footer.add_table(1, 3, Cm(18))
        tbl_f.autofit = False
        tbl_f.allow_autofit = False
        
        for row in tbl_f.rows:
            for idx, width in enumerate(header_widths):
                row.cells[idx].width = width
        # Izquierda: centro (en blanco → rellenar)
        c0 = tbl_f.rows[0].cells[0]
        c0.paragraphs[0].add_run("IES / Centro").font.size = Pt(8)
        # Centro: vacío
        c1 = tbl_f.rows[0].cells[1]
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Derecha: número de página
        c2 = tbl_f.rows[0].cells[2]
        p_pag = c2.paragraphs[0]
        p_pag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_lbl = p_pag.add_run("Página: ")
        run_lbl.font.size = Pt(8)
        # Campo de número de página automático
        run_pg = p_pag.add_run()
        fldChar1 = OxmlElement("w:fldChar"); fldChar1.set(qn("w:fldCharType"), "begin")
        instrT   = OxmlElement("w:instrText"); instrT.set(qn("xml:space"), "preserve"); instrT.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar"); fldChar2.set(qn("w:fldCharType"), "end")
        run_pg._r.extend([fldChar1, instrT, fldChar2])
        run_pg.font.size = Pt(8)


# ─── PORTADA ────────────────────────────────────────────────────────────────

def build_cover(doc, data):
    """Portada idéntica en estructura al PDF de referencia"""
    doc.add_paragraph()  # espacio superior
    doc.add_paragraph()
    doc.add_paragraph()

    def cover_field(label, value):
        p_lbl = doc.add_paragraph()
        p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_lbl = p_lbl.add_run(f"{label}:")
        set_font(r_lbl, bold=True, size=12)
        set_para_spacing(p_lbl, before=12, after=2)

        p_val = doc.add_paragraph()
        p_val.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_val = p_val.add_run(value)
        set_font(r_val, size=12)
        set_para_spacing(p_val, before=2, after=12)

    # Título
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("PROGRAMACIÓN DIDÁCTICA")
    set_font(r_title, bold=True, size=20, color=(0, 0, 0))
    set_para_spacing(p_title, before=20, after=30)

    cover_field("DEPARTAMENTO DIDÁCTICO", data.get("departamento", "Electrónica"))
    cover_field("CICLO FORMATIVO", data.get("ciclo", "Ciclo formativo"))
    cover_field("MÓDULO PROFESIONAL", data.get("modulo", "Módulo profesional"))
    cover_field("CURSO ACADÉMICO", data.get("curso_academico", "2025/2026"))

    doc.add_page_break()


# ─── SECCIONES ──────────────────────────────────────────────────────────────

def build_sections(doc, data):
    ras  = data.get("df_ra", [])
    uds  = data.get("df_ud", [])
    acts = data.get("df_act", [])
    ces  = data.get("df_ce", [])
    info = data.get("info_modulo", {})
    df_sgmt = data.get("df_sgmt", [])
    df_sesiones = data.get("df_sesiones", [])

    curriculo_data = data.get("curriculo_data", {})
    boa_articles = curriculo_data.get("boa_articles", {})
    competencias_cpps = curriculo_data.get("competencias_cpps", [])
    objetivos_generales = curriculo_data.get("objetivos_generales_og", [])
    unidades_formativas = curriculo_data.get("unidades_formativas", [])

    modulo_nombre = data.get("modulo", info.get("modulo", "Módulo"))
    horas_totales = data.get("horas_totales", info.get("horas", "N/D"))
    ctx = data.get("config_contexto", {})

    def add_context_body(doc, key, placeholder):
        val = ctx.get(key, "").strip()
        if val:
            add_body(doc, val)
        else:
            p = add_body(doc, f"[Pendiente de redactar por el departamento: {placeholder}]")
            p.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    # ──────────────────────────────────────────────────────────
    # A. INTRODUCCIÓN Y CONTEXTUALIZACIÓN
    add_heading(doc, "A. Introducción y contextualización", 1)
    
    add_heading(doc, "A1. Justificación de la programación", 2)
    add_context_body(doc, "A1_justificacion", "Indicar base normativa, Leyes de Educación y Reales Decretos aplicables al título...")

    add_heading(doc, "A2. Contextualización", 2)
    add_context_body(doc, "A2_contextualizacion", "Perfil profesional del título, entorno socioeconómico y características generales del centro...")

    add_heading(doc, "A3. Ubicación del módulo en el ciclo formativo", 2)
    add_body(doc,
        f"El presente módulo profesional, '{modulo_nombre}', se imparte dentro del ciclo formativo "
        f"y tiene una carga horaria de {horas_totales} horas. "
        "Aporta al alumnado los conocimientos, procedimientos y actitudes necesarios para el ejercicio "
        "profesional en el sector correspondiente, en coherencia con los Resultados de Aprendizaje."
    )
    if unidades_formativas:
        add_body(doc, "Se estructura en las siguientes Unidades Formativas (UF):")
        for uf in unidades_formativas:
            add_bullet(doc, f"{uf.get('codigo', '')}: {uf.get('nombre', '')} ({uf.get('horas', 0)}h)")

    add_heading(doc, "A4. Identificación del título y perfil profesional", 2)
    art_keys = ["article_2", "article_3", "article_4", "article_6", "article_7", "article_8"]
    for ak in art_keys:
        val = boa_articles.get(ak)
        if val:
            lines = val.split('\n')
            if lines:
                add_heading(doc, lines[0].strip(), 3)
                for line in lines[1:]:
                    if line.strip():
                        add_body(doc, line.strip())

    if competencias_cpps:
        add_heading(doc, "Competencias profesionales, personales y sociales", 3)
        add_body(doc, "La contribución del módulo a las competencias profesionales, personales y sociales del título es la siguiente:")
        for cpp in competencias_cpps:
            add_bullet(doc, f"CPPS {cpp.get('id', '')}. {cpp.get('descripcion', '')}")

    # ──────────────────────────────────────────────────────────
    # B. RESULTADOS DE APRENDIZAJE Y FCT/FEOE
    add_heading(doc, "B. Resultados de aprendizaje y FCT/FEOE", 1)
    
    add_heading(doc, "B1. Resultados de aprendizaje del módulo", 2)
    add_body(doc, f"Esta programación didáctica consta de {len(ras)} Resultado(s) de Aprendizaje oficiales:")
    for ra in ras:
        add_bullet(doc, f"{ra.get('id_ra', '')}. {ra.get('desc_ra', '')}")

    add_heading(doc, "B2. RAs susceptibles de ser adquiridos en FEOE", 2)
    feoe_ras = [ra for ra in ras if ra.get("is_dual")]
    if feoe_ras:
        add_body(doc, "Los siguientes RAs son susceptibles de ser trabajados en fase de empresa (FEOE):")
        for ra in feoe_ras:
            add_bullet(doc, f"{ra.get('id_ra', '')}. {ra.get('desc_ra', '')}")
    else:
        add_body(doc, "No se han marcado RAs específicos para Formación en Empresa. Todos son susceptibles a criterio docente.")

    add_heading(doc, "B3. Vinculación con la empresa colaboradora", 2)
    add_context_body(doc, "B3_vinculacion_empresa", "Orientaciones sobre las actividades a realizar en la empresa...")

    if objetivos_generales:
        add_heading(doc, "B4. Objetivos generales del título y su contribución", 2)
        add_body(doc, "Los objetivos generales del título a los que contribuye este módulo son los siguientes:")
        for og in objetivos_generales:
            add_bullet(doc, f"OG {og.get('id', '')}. {og.get('descripcion', '')}")

    # ──────────────────────────────────────────────────────────
    # C. ORGANIZACIÓN, SECUENCIACIÓN Y TEMPORALIZACIÓN
    add_heading(doc, "C. Organización, secuenciación y temporalización", 1)
    
    add_heading(doc, "C1. Unidades didácticas (contenidos)", 2)
    if uds:
        tbl = doc.add_table(rows=1, cols=3)
        tbl.style = "Table Grid"
        tbl.autofit = False
        tbl.allow_autofit = False
        
        widths = [Cm(2), Cm(14), Cm(2)]
        for i, width in enumerate(widths):
            tbl.columns[i].width = width

        for i, txt in enumerate(["UD", "UD", "Horas"]):
            cell = tbl.rows[0].cells[i]
            cell.width = widths[i]
            r = cell.paragraphs[0].add_run(txt)
            set_font(r, bold=True, size=10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        for ud in uds:
            row = tbl.add_row().cells
            ud_id = str(ud.get("id_ud", ""))
            
            for i in range(3):
                row[i].width = widths[i]
                
            row[0].paragraphs[0].add_run(ud_id).font.size = Pt(9)
            row[1].paragraphs[0].add_run(str(ud.get("desc_ud", ""))).font.size = Pt(9)
            row[2].paragraphs[0].add_run(str(ud.get("horas_ud", ""))).font.size = Pt(9)
            row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        add_body(doc, "[Pendiente de definir las Unidades Didácticas]")

    add_heading(doc, "C2. Secuenciación y distribución trimestral", 2)
    add_body(doc, "Distribución estimada de las UDs a lo largo del curso (Planificación mensual):")
    if df_sgmt:
        # Generate table for planning
        months = ["Sep", "Oct", "Nov", "Dic", "Ene", "Feb", "Mar", "Abr", "May", "Jun"]
        tbl2 = doc.add_table(rows=1, cols=len(months) + 2)
        tbl2.style = "Table Grid"
        tbl2.rows[0].cells[0].paragraphs[0].add_run("UD").bold = True
        tbl2.rows[0].cells[1].paragraphs[0].add_run("Horas").bold = True
        for i, m in enumerate(months):
            tbl2.rows[0].cells[i+2].paragraphs[0].add_run(m).bold = True
        for row_data in df_sgmt:
            tr = tbl2.add_row().cells
            tr[0].paragraphs[0].add_run(str(row_data.get("id_ud", ""))).font.size = Pt(8)
            tr[1].paragraphs[0].add_run(str(row_data.get("horas_ud", ""))).font.size = Pt(8)
            for i, m in enumerate(months):
                val = row_data.get(f"{m}_Prv", "")
                if val: tr[i+2].paragraphs[0].add_run(str(val)).font.size = Pt(8)
    else:
        p_c2 = add_body(doc, "[Pendiente de auto-planificar el calendario en CuadernoFP]")
        p_c2.runs[0].font.color.rgb = RGBColor(128, 128, 128)

    add_heading(doc, "C3. Calendario académico y carga horaria", 2)
    add_body(doc, "La secuenciación se ajusta a los festivos oficiales y al horario semanal establecido para el grupo.")

    add_heading(doc, "C4. Matrices de relación RA - UD", 2)
    if not ras or not uds:
        add_body(doc, "[Pendiente de definir RAs o UDs para generar la matriz]")
    else:
        ra_ids = [ra.get("id_ra") for ra in ras if ra.get("id_ra")]
        tbl3 = doc.add_table(rows=1, cols=2 + len(ra_ids))
        tbl3.style = "Table Grid"
        
        # Header
        hcells = tbl3.rows[0].cells
        r_ud = hcells[0].paragraphs[0].add_run("UD")
        set_font(r_ud, bold=True, size=9)
        hcells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_hr = hcells[1].paragraphs[0].add_run("Horas")
        set_font(r_hr, bold=True, size=9)
        hcells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for i, ra_id in enumerate(ra_ids):
            r_ra = hcells[2+i].paragraphs[0].add_run(ra_id)
            set_font(r_ra, bold=True, size=9)
            hcells[2+i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        for ud in uds:
            row_cells = tbl3.add_row().cells
            row_cells[0].paragraphs[0].add_run(str(ud.get("id_ud", ""))).font.size = Pt(8)
            row_cells[1].paragraphs[0].add_run(str(ud.get("horas_ud", ""))).font.size = Pt(8)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for i, ra_id in enumerate(ra_ids):
                val = ud.get(ra_id, 0)
                try:
                    val_float = float(val or 0)
                except ValueError:
                    val_float = 0
                txt = f"{int(val_float)}%" if val_float > 0 else ""
                r_val = row_cells[2+i].paragraphs[0].add_run(txt)
                r_val.font.size = Pt(8)
                if val_float > 0: r_val.bold = True
                row_cells[2+i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ──────────────────────────────────────────────────────────
    # D. METODOLOGÍA DIDÁCTICA Y AGRUPAMIENTOS
    add_heading(doc, "D. Metodología didáctica y agrupamientos", 1)
    
    add_heading(doc, "D1. Principios metodológicos generales", 2)
    add_body(doc, "La metodología favorece en el alumnado la capacidad para aprender de forma autónoma y trabajar en equipo.")
    
    add_heading(doc, "D2. Actividades de enseñanza-aprendizaje", 2)
    add_context_body(doc, "D2_actividades_ea", "Relación de metodologías tipo como teoría, taller, prácticas simuladas...")

    add_heading(doc, "D3. Agrupamientos y plan de desdobles", 2)
    add_context_body(doc, "D3_agrupamientos", "Organización del grupo, desdobles por prevención de riesgos o ratios...")

    # ──────────────────────────────────────────────────────────
    # E. PROCEDIMIENTOS, CRITERIOS E INSTRUMENTOS DE EVALUACIÓN
    add_heading(doc, "E. Procedimientos, criterios e instrumentos de evaluación", 1)
    
    add_heading(doc, "E1. Evaluación inicial", 2)
    add_body(doc, "Al inicio del módulo se realizará una evaluación diagnóstica para identificar conocimientos previos.")

    add_heading(doc, "E2. Criterios de evaluación", 2)
    if ras:
        for ra in ras:
            p = doc.add_paragraph()
            set_para_spacing(p, before=12, after=2)
            r = p.add_run(f"{ra.get('id_ra', '')}")
            set_font(r, bold=True, size=10)
            ra_ces = [ce for ce in ces if ce.get("id_ra") == ra.get("id_ra")]
            for ce in ra_ces:
                add_bullet(doc, f"{ce.get('id_ce', '')}. {ce.get('desc_ce', '')} ({ce.get('peso_ce', 0)}%)")
    else:
        add_body(doc, "[Pendiente de definir los RA y CE]")

    add_heading(doc, "E3. Procedimientos e instrumentos de evaluación", 2)
    if acts:
        tipos = {}
        for act in acts:
            t = act.get("Tipo", "Otro")
            tipos[t] = tipos.get(t, 0) + 1
        add_body(doc, "Se utilizarán los siguientes instrumentos de evaluación configurados en la plataforma:")
        for tipo, n in tipos.items():
            add_bullet(doc, f"{tipo}: {n} actividad(es)")
    else:
        add_body(doc, "[Pendiente de crear actividades]")

    add_heading(doc, "E4. Criterios de calificación y ponderación", 2)
    config = data.get("config_redondeo", {})
    nota_ap = config.get("nota_aprobado", 5.0)
    umbral  = config.get("umbral_redondeo", 4.5)
    add_body(doc, f"El alumnado debe alcanzar un mínimo de {nota_ap} en cada Resultado de Aprendizaje. "
                  f"Las calificaciones entre {umbral} y {nota_ap} podrán ser objeto de redondeo.")

    add_heading(doc, "E5. Actividades de recuperación y refuerzo", 2)
    add_context_body(doc, "E5_recuperacion", "Sistema de recuperación para evaluaciones y convocatorias extraordinarias...")

    # ──────────────────────────────────────────────────────────
    # F. ATENCIÓN A LAS DIFERENCIAS INDIVIDUALES
    add_heading(doc, "F. Atención a las diferencias individuales", 1)
    
    add_heading(doc, "F1. Atención a la diversidad", 2)
    add_context_body(doc, "F1_diversidad", "Medidas de inclusión y atención a las diferencias individuales...")

    add_heading(doc, "F2. Adaptaciones curriculares", 2)
    add_body(doc, "Se contemplarán adaptaciones no significativas de acceso al currículo según lo requerido.")

    # ──────────────────────────────────────────────────────────
    # G. MATERIALES Y RECURSOS DIDÁCTICOS
    add_heading(doc, "G. Materiales y recursos didácticos", 1)
    
    add_heading(doc, "G1. Infraestructuras y equipamientos", 2)
    add_context_body(doc, "G1_infraestructuras", "Taller, aula, laboratorio...")

    add_heading(doc, "G2. Herramientas TIC y plataformas", 2)
    add_context_body(doc, "G2_herramientas_tic", "Moodle, Classroom, software simulador...")

    add_heading(doc, "G3. Bibliografía y recursos para el alumnado", 2)
    add_context_body(doc, "G3_bibliografia", "Libros de texto, manuales de fabricantes...")

    # ──────────────────────────────────────────────────────────
    # H. ACTIVIDADES COMPLEMENTARIAS Y EXTRAESCOLARES
    add_heading(doc, "H. Actividades complementarias y extraescolares", 1)
    add_heading(doc, "H1. Propuestas del departamento", 2)
    add_context_body(doc, "H1_complementarias", "Actividades extraescolares y complementarias propuestas...")

    # ──────────────────────────────────────────────────────────
    # I. ELEMENTOS TRANSVERSALES Y PROYECTOS ESPECÍFICOS
    add_heading(doc, "I. Elementos transversales y proyectos específicos", 1)
    add_heading(doc, "I1. Elementos transversales", 2)
    add_context_body(doc, "I1_transversales", "Prevención de riesgos laborales, igualdad, medioambiente...")

    add_heading(doc, "I2. Proyectos de innovación o bilingüismo", 2)
    add_body(doc, "La programación no contempla modalidad bilingüe en el presente curso académico, salvo indicación contraria.")

    # ──────────────────────────────────────────────────────────
    # J. SEGUIMIENTO Y EVALUACIÓN DE LA PROGRAMACIÓN DIDÁCTICA
    add_heading(doc, "J. Seguimiento y evaluación de la programación didáctica", 1)
    
    add_heading(doc, "J1. Mecanismos de seguimiento docente", 2)
    add_body(doc, "El Departamento revisará periódicamente el cumplimiento de la programación y analizará la evolución de resultados.")

    add_heading(doc, "J2. Indicadores de logro", 2)
    add_body(doc, "El control del avance de las Unidades Didácticas y la docencia diaria se registrará a través de la herramienta digital CuadernoFP. El diario docente telemático servirá como indicador objetivo del grado de cumplimiento temporal respecto al calendario planificado.")

    add_heading(doc, "J3. Plan de contingencia", 2)
    add_context_body(doc, "J3_contingencia", "Procedimiento ante ausencias o clases a distancia...")

    add_heading(doc, "J4. Registro de seguimiento diario (diario docente)", 2)
    if not df_sesiones:
        add_body(doc, "No hay sesiones registradas en el diario docente.")
    else:
        add_body(doc, "A continuación se detalla el registro de sesiones impartidas, obtenido automáticamente del seguimiento diario:")
        
        tbl_sg = doc.add_table(rows=1, cols=4)
        tbl_sg.style = "Table Grid"
        headers_sg = ["Fecha", "Tipo", "UD", "Contenido Impartido"]
        for i, txt in enumerate(headers_sg):
            r = tbl_sg.rows[0].cells[i].paragraphs[0].add_run(txt)
            set_font(r, bold=True, size=9)
            tbl_sg.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Ajustar ancho de columnas si es posible
        from docx.shared import Cm
        tbl_sg.columns[0].width = Cm(2.5)
        tbl_sg.columns[1].width = Cm(2.5)
        tbl_sg.columns[2].width = Cm(2)
        tbl_sg.columns[3].width = Cm(10)
        
        for ses in sorted(df_sesiones, key=lambda x: str(x.get("fecha", ""))):
            row_cells = tbl_sg.add_row().cells
            
            fecha = str(ses.get("fecha", ""))
            if len(fecha) >= 10 and "-" in fecha:
                parts = fecha[:10].split("-")
                if len(parts) == 3:
                    fecha = f"{parts[2]}/{parts[1]}/{parts[0]}"
                    
            row_cells[0].paragraphs[0].add_run(fecha).font.size = Pt(8)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_cells[1].paragraphs[0].add_run(str(ses.get("tipo_sesion", ""))).font.size = Pt(8)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            row_cells[2].paragraphs[0].add_run(str(ses.get("ud_id", ""))).font.size = Pt(8)
            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            tit = str(ses.get("titulo_sesion", ""))
            obs = str(ses.get("observaciones", "")).strip()
            if obs:
                tit += f" | Obs: {obs}"
            row_cells[3].paragraphs[0].add_run(tit).font.size = Pt(8)


# ─── MAIN ───────────────────────────────────────────────────────────────────

def generate(data, output_path_docx, output_path_pdf=None):
    """
    data: dict con claves:
        departamento, ciclo, modulo, curso_academico, horas_totales,
        df_ra, df_ud, df_act, df_ce, config_redondeo, info_modulo
    """
    doc = Document()

    # Estilos de párrafo Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0, 0, 0)

    # Configurar encabezado y pie
    modulo_nombre    = data.get("modulo", data.get("info_modulo", {}).get("modulo", "Módulo"))
    ciclo_nombre     = data.get("ciclo", data.get("info_modulo", {}).get("ciclo", "Ciclo Formativo"))
    curso_academico  = data.get("curso_academico", "2025/2026")
    setup_header_footer(doc, modulo_nombre, ciclo_nombre, curso_academico)

    # Portada
    build_cover(doc, data)

    # Índice automático (se actualiza al abrir en Word)
    add_heading(doc, "Índice", 1)
    doc.add_paragraph()
    add_toc(doc)
    doc.add_page_break()

    # Cuerpo del documento
    build_sections(doc, data)

    # Guardar DOCX
    doc.save(output_path_docx)
    print(f"[OK] DOCX generado: {output_path_docx}")

    if output_path_pdf:
        import subprocess
        out_dir = os.path.dirname(output_path_pdf)
        try:
            result = subprocess.run(
                ["soffice", "--headless", "--convert-to", "pdf",
                 "--outdir", out_dir, output_path_docx],
                capture_output=True, text=True
            )
            if result.returncode == 0:
                # LibreOffice usa el mismo nombre base con extensión .pdf
                base = os.path.splitext(os.path.basename(output_path_docx))[0]
                generated_pdf = os.path.join(out_dir, base + ".pdf")
                if generated_pdf != output_path_pdf and os.path.exists(generated_pdf):
                    os.rename(generated_pdf, output_path_pdf)
                print(f"[OK] PDF generado: {output_path_pdf}")
            else:
                print(f"[WARN] LibreOffice devolvió error al intentar generar PDF.")
                print(f"  Abre el DOCX en Word y exporta como PDF desde alli.")
        except FileNotFoundError:
            print(f"[WARN] LibreOffice no disponible. Instala LibreOffice para generar PDF automaticamente.")
            print(f"  Abre el DOCX en Word y exporta como PDF desde alli.")


# ─── Demo / CLI ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    # Datos de demostración (se sustituyen por los reales al llamar desde el backend)
    demo_data = {
        "departamento": "Electrónica",
        "ciclo": "Técnico en Instalaciones de Telecomunicaciones",
        "modulo": "0237 Infraestructuras comunes de telecomunicación en viviendas y edificios",
        "curso_academico": "2025/2026",
        "horas_totales": 224,
        "df_ra": [
            {"id_ra": "RA1", "desc_ra": "Identifica los elementos de las ICT...", "is_dual": False},
            {"id_ra": "RA2", "desc_ra": "Configura pequeñas instalaciones ICT...", "is_dual": True},
        ],
        "df_ud": [
            {"id_ud": "UD1", "desc_ud": "Introducción a las ICT", "horas_ud": 20},
            {"id_ud": "UD2", "desc_ud": "Instalaciones de antenas", "horas_ud": 30},
        ],
        "df_act": [
            {"Tipo": "Teoria", "id_act": "ACT01"},
            {"Tipo": "Practica", "id_act": "ACT02"},
            {"Tipo": "Informes", "id_act": "ACT03"},
        ],
        "df_ce": [
            {"id_ce": "CE1.1", "desc_ce": "Se han identificado los sistemas ICT.", "id_ra": "RA1"},
            {"id_ce": "CE1.2", "desc_ce": "Se han descrito los tipos de señal.", "id_ra": "RA1"},
        ],
        "config_redondeo": {"nota_aprobado": 5.0, "umbral_redondeo": 4.5, "max_compensables": 0},
        "info_modulo": {}
    }

    out_dir = sys.argv[2] if len(sys.argv) > 2 else "."
    fname   = f"PD_{demo_data['modulo'][:30].replace(' ', '_').replace('/', '-')}"
    out_docx = os.path.join(out_dir, fname + ".docx")
    out_pdf  = os.path.join(out_dir, fname + ".pdf")

    generate(demo_data, out_docx, out_pdf)
